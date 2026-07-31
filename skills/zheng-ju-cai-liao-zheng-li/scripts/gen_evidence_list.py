# -*- coding: utf-8 -*-
"""
证据清单 Word 文档生成器
用法:
    python gen_evidence_list.py <输出路径> <JSON证据数据>
示例:
    python gen_evidence_list.py "D:\\输出\\证据清单.docx" "[{\"num\":\"1.1\",\"name\":\"收条\",\"content\":\"证明：...\",\"pages\":\"第1-2页\"}]"
"""
import sys, json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

PYE = r'C:\Program Files\QClaw\v0.2.29.592\resources\python\python.exe'
FONT = '宋体'
FONT_HINT = '仿宋_GB2312'

def set_cell_text(cell, text, bold=False, size=12, font_name=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name or FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or FONT)

def add_para(doc, text, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, font_name=None):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name or FONT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or FONT)
    return p

def set_row_height(row, height_pt):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

def set_col_width(table, col_idx, width_dxa):
    for row in table.rows:
        if col_idx < len(row.cells):
            row.cells[col_idx].width = width_dxa

def add_footer(doc, total_pages):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    for para in footer.paragraphs:
        for run in para.runs:
            run.clear()
    if not footer.paragraphs:
        footer.add_paragraph()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run('第 ')
    run.font.name = FONT_HINT
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HINT)
    run.font.size = Pt(12)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    fldChar2 = OxmlElement('w:instrText')
    fldChar2.text = ' PAGE '
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run2 = fp.add_run()
    run2.font.name = FONT_HINT
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HINT)
    run2.font.size = Pt(12)
    run2._r.append(fldChar1)
    run2._r.append(fldChar2)
    run2._r.append(fldChar3)
    run3 = fp.add_run(' 页 共 ')
    run3.font.name = FONT_HINT
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HINT)
    run3.font.size = Pt(12)
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'begin')
    fldChar5 = OxmlElement('w:instrText')
    fldChar5.text = ' NUMPAGES '
    fldChar6 = OxmlElement('w:fldChar')
    fldChar6.set(qn('w:fldCharType'), 'end')
    run4 = fp.add_run()
    run4.font.name = FONT_HINT
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HINT)
    run4.font.size = Pt(12)
    run4._r.append(fldChar4)
    run4._r.append(fldChar5)
    run4._r.append(fldChar6)
    run5 = fp.add_run(' 页')
    run5.font.name = FONT_HINT
    run5._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HINT)
    run5.font.size = Pt(12)

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_table_width(table, width_dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def set_cell_vertical_align(cell, align='center'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def build_doc(output_path, evidence_data, title, signatory='提交人'):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # 标题
    add_para(doc, title, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 表格
    # 列宽(合计=12480DXA, A4宽210mm-50mm边距)
    col_widths = [1200, 3600, 5280, 2400]  # 编号, 证据名称, 证明内容, 页码
    total_w = sum(col_widths)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(table)
    set_table_width(table, total_w)

    # 表头
    header_row = table.rows[0]
    headers = ['编号', '证据名称', '证据主要内容及证明对象', '页码']
    for ci, hdr in enumerate(headers):
        cell = header_row.cells[ci]
        cell.text = ''
        set_cell_margins(cell)
        set_col_width(table, ci, col_widths[ci])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        set_cell_vertical_align(cell)

    # 数据行
    for item in evidence_data:
        row = table.add_row()
        cells = row.cells

        # 编号
        set_cell_margins(cells[0])
        set_col_width(table, 0, col_widths[0])
        p0 = cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = p0.add_run(item['num'])
        run0.font.size = Pt(12)
        run0.font.name = FONT
        run0._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        set_cell_vertical_align(cells[0])

        # 证据名称
        set_cell_margins(cells[1])
        set_col_width(table, 1, col_widths[1])
        p1 = cells[1].paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run1 = p1.add_run(item['name'])
        run1.font.size = Pt(12)
        run1.font.name = FONT
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        set_cell_vertical_align(cells[1])

        # 证明内容
        set_cell_margins(cells[2])
        set_col_width(table, 2, col_widths[2])
        p2 = cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run2 = p2.add_run(item['content'])
        run2.font.size = Pt(12)
        run2.font.name = FONT
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        set_cell_vertical_align(cells[2])

        # 页码
        set_cell_margins(cells[3])
        set_col_width(table, 3, col_widths[3])
        p3 = cells[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(item['pages'])
        run3.font.size = Pt(12)
        run3.font.name = FONT
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        set_cell_vertical_align(cells[3])

    doc.add_paragraph()
    # 落款
    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_s = p_sign.add_run(f'{signatory}（原告）：______________    日期：_____年_____月_____日')
    run_s.font.size = Pt(12)
    run_s.font.name = FONT_HINT
    run_s._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HINT)

    # 页脚
    add_footer(doc, 1)

    doc.save(output_path)
    print(f'SAVED: {output_path}')

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    output = sys.argv[1]
    data = json.loads(sys.argv[2])
    title = sys.argv[3] if len(sys.argv) > 3 else '证据目录（原告提交）'
    signatory = sys.argv[4] if len(sys.argv) > 4 else '提交人'
    build_doc(output, data, title, signatory)
