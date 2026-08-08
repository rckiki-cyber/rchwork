#!/usr/bin/env python3
"""Semantic legal-document DOCX formatter for LegalWork.

This worker captures recurring formatting conventions from real legal-work
samples without pushing paragraph-by-paragraph Office commands into model
history. Chinese body text defaults to SimSun (宋体) 12 pt (小四).
"""
from __future__ import annotations

import argparse
import json
import shutil
import zipfile
from collections import Counter
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except Exception as exc:  # pragma: no cover
    Document = None  # type: ignore[assignment]
    IMPORT_ERROR = str(exc)
else:
    IMPORT_ERROR = ""

COMMON_PAGE = {
    "width_cm": 21.0,
    "height_cm": 29.7,
    "top_cm": 2.54,
    "bottom_cm": 2.54,
    "left_cm": 3.17,
    "right_cm": 3.17,
    "header_cm": 1.5,
    "footer_cm": 1.75,
}

# All profiles intentionally keep the user's default Chinese body font at
# 宋体、小四. Differences are structural: spacing, indentation and table use.
PROFILES: dict[str, dict[str, Any]] = {
    "fact-memo": {
        "body": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "line_spacing": 1.5, "first_indent_pt": 24.0, "alignment": "justify"},
        "headings": {
            "Heading 1": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True, "first_indent_pt": 0.0},
            "Heading 2": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True, "first_indent_pt": 0.0},
        },
        "page": COMMON_PAGE,
        "safe_default_scopes": ["page", "body", "headings"],
    },
    "legal-research": {
        "body": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "line_spacing": 1.5, "first_indent_pt": 24.0, "alignment": "justify"},
        "headings": {
            "Heading 3": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True, "first_indent_pt": 0.0},
            "Heading 4": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True, "first_indent_pt": 0.0},
        },
        "page": COMMON_PAGE,
        "safe_default_scopes": ["page", "body", "headings"],
    },
    "engagement-agreement": {
        "body": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "line_spacing": 1.25, "first_indent_pt": 0.0, "alignment": "left"},
        "headings": {
            "Title": {"east_asia": "宋体", "latin": "Times New Roman", "size": 14.0, "bold": True, "first_indent_pt": 0.0, "alignment": "center"},
        },
        "page": COMMON_PAGE,
        "safe_default_scopes": ["page", "body"],
    },
    "case-notes": {
        "body": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "line_spacing": 1.5, "first_indent_pt": 0.0, "alignment": "justify"},
        "headings": {
            "Heading 1": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True, "first_indent_pt": 0.0},
            "Heading 2": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True, "first_indent_pt": 0.0},
            "Heading 4": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True, "first_indent_pt": 0.0},
        },
        "page": COMMON_PAGE,
        # Mixed case notes often contain comparison tables and intentional local
        # formatting. Do not flatten table contents unless explicitly requested.
        "safe_default_scopes": ["page", "body", "headings"],
    },
    "case-tables": {
        "body": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "line_spacing": 1.25, "first_indent_pt": 0.0, "alignment": "left"},
        "headings": {},
        "table": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "line_spacing": 1.25, "first_indent_pt": 0.0},
        "page": COMMON_PAGE,
        "safe_default_scopes": ["page", "tables"],
    },
}

SPECIAL_STYLE_PREFIXES = (
    "title", "subtitle", "heading", "caption", "quote", "footnote",
    "header", "footer", "toc", "目录", "标题", "题注", "引文",
)


def emit(payload: dict[str, Any], code: int = 0) -> None:
    print(json.dumps(payload, ensure_ascii=False, separators=(",", ":")))
    raise SystemExit(code)


def fail(message: str, operation: str, code: int = 1) -> None:
    emit({"status": "error", "operation": operation, "error": message}, code)


def require_docx(operation: str) -> None:
    if Document is None:
        fail(f"python-docx unavailable: {IMPORT_ERROR}", operation)


def input_docx(raw: str, operation: str) -> Path:
    path = Path(raw).expanduser().resolve()
    if not path.is_file():
        fail(f"file not found: {path}", operation)
    if path.suffix.lower() != ".docx":
        emit({"status": "unsupported", "operation": operation, "reason": "semantic legal profiles require .docx input", "office_fallback_allowed": False})
    return path


def output_docx(raw: str, operation: str) -> Path:
    path = Path(raw).expanduser().resolve()
    if path.suffix.lower() != ".docx":
        fail("output must end with .docx", operation)
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def risky_features(path: Path) -> dict[str, bool]:
    result = {"tracked_changes": False, "macros": False}
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
            result["macros"] = any(name.lower().endswith("vbaproject.bin") for name in names)
            for name in ("word/document.xml", "word/footnotes.xml", "word/endnotes.xml"):
                if name not in names:
                    continue
                xml = zf.read(name).decode("utf-8", errors="ignore")
                if any(token in xml for token in ("<w:ins", "<w:del", "<w:moveFrom", "<w:moveTo")):
                    result["tracked_changes"] = True
                    break
    except zipfile.BadZipFile:
        fail("invalid DOCX zip container", "profile-apply")
    return result


def set_font(target: Any, spec: dict[str, Any], *, force_black: bool = False) -> None:
    latin = spec.get("latin") or "Times New Roman"
    east_asia = spec.get("east_asia") or "宋体"
    size = float(spec.get("size") or 12.0)
    target.font.name = latin
    target.font.size = Pt(size)
    if force_black:
        target.font.color.rgb = RGBColor(0, 0, 0)
    if spec.get("bold") is not None:
        target.font.bold = bool(spec["bold"])
    fonts = target._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), latin)
    fonts.set(qn("w:hAnsi"), latin)
    fonts.set(qn("w:cs"), latin)
    fonts.set(qn("w:eastAsia"), east_asia)
    fonts.set(qn("w:hint"), "eastAsia")


def remove_xml_children(parent: Any, *qualified_names: str) -> None:
    if parent is None:
        return
    wanted = {qn(name) for name in qualified_names}
    for child in list(parent):
        if child.tag in wanted:
            parent.remove(child)


def remove_builtin_heading_residue(style: Any) -> None:
    ppr = style._element.get_or_add_pPr()
    rpr = style._element.get_or_add_rPr()
    remove_xml_children(ppr, "w:pBdr", "w:shd", "w:tabs")
    remove_xml_children(rpr, "w:spacing", "w:shd", "w:effect", "w:outline", "w:shadow")


def set_paragraph_format(target: Any, spec: dict[str, Any]) -> None:
    fmt = target.paragraph_format
    if spec.get("line_spacing") is not None:
        fmt.line_spacing = float(spec["line_spacing"])
    if spec.get("first_indent_pt") is not None:
        fmt.first_line_indent = Pt(float(spec["first_indent_pt"]))
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    alignment = spec.get("alignment")
    if alignment == "justify":
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == "center":
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "left":
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT


def apply_page(doc: Any, page: dict[str, float]) -> int:
    for section in doc.sections:
        section.page_width = Cm(page["width_cm"])
        section.page_height = Cm(page["height_cm"])
        section.top_margin = Cm(page["top_cm"])
        section.bottom_margin = Cm(page["bottom_cm"])
        section.left_margin = Cm(page["left_cm"])
        section.right_margin = Cm(page["right_cm"])
        section.header_distance = Cm(page["header_cm"])
        section.footer_distance = Cm(page["footer_cm"])
    return len(doc.sections)


def is_body_paragraph(paragraph: Any) -> bool:
    if not paragraph.text.strip():
        return False
    name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
    return not any(name.startswith(prefix) for prefix in SPECIAL_STYLE_PREFIXES)


def apply_body(doc: Any, spec: dict[str, Any]) -> dict[str, int]:
    changed = Counter()
    try:
        normal = doc.styles["Normal"]
        set_font(normal, spec)
        set_paragraph_format(normal, spec)
        changed["styles"] += 1
    except Exception:
        pass
    for paragraph in doc.paragraphs:
        if not is_body_paragraph(paragraph):
            continue
        set_paragraph_format(paragraph, spec)
        for run in paragraph.runs:
            if run.text:
                set_font(run, spec)
                changed["runs"] += 1
        changed["paragraphs"] += 1
    return dict(changed)


def apply_headings(doc: Any, specs: dict[str, dict[str, Any]]) -> dict[str, int]:
    changed = Counter()
    for style_name, spec in specs.items():
        try:
            style = doc.styles[style_name]
            set_font(style, spec, force_black=True)
            set_paragraph_format(style, spec)
            remove_builtin_heading_residue(style)
            changed["styles"] += 1
        except KeyError:
            continue
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        spec = specs.get(style_name)
        if not spec:
            continue
        set_paragraph_format(paragraph, spec)
        for run in paragraph.runs:
            if run.text:
                set_font(run, spec, force_black=True)
                changed["runs"] += 1
        changed["paragraphs"] += 1
    return dict(changed)


def apply_tables(doc: Any, spec: dict[str, Any]) -> dict[str, int]:
    changed = Counter()
    for table in doc.tables:
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    set_paragraph_format(paragraph, spec)
                    for run in paragraph.runs:
                        if run.text:
                            set_font(run, {**spec, "bold": True if row_index == 0 else run.font.bold})
                            changed["runs"] += 1
                    changed["paragraphs"] += 1
        changed["tables"] += 1
    return dict(changed)


def apply_structure_spec(doc: Any, spec_path: str) -> dict[str, int]:
    payload = json.loads(Path(spec_path).expanduser().resolve().read_text(encoding="utf-8"))
    assignments = payload.get("styles") if isinstance(payload, dict) else None
    if not isinstance(assignments, list):
        fail("structure spec must contain a styles array", "structure")
    changed = Counter()
    for entry in assignments:
        if not isinstance(entry, dict):
            continue
        style_name = str(entry.get("style") or "").strip()
        if style_name not in {"Title", "Heading 1", "Heading 2", "Heading 3", "Heading 4", "Normal"}:
            continue
        indexes = entry.get("paragraph_indexes")
        exact = entry.get("exact_text")
        candidates: list[Any] = []
        if isinstance(indexes, list):
            for raw in indexes:
                try:
                    idx = int(raw)
                except Exception:
                    continue
                if 0 <= idx < len(doc.paragraphs):
                    candidates.append(doc.paragraphs[idx])
        if isinstance(exact, str) and exact:
            candidates.extend(p for p in doc.paragraphs if p.text.strip() == exact.strip())
        for paragraph in candidates:
            paragraph.style = doc.styles[style_name]
            changed[style_name] += 1
    return dict(changed)


def merge_counts(target: Counter[str], prefix: str, values: dict[str, int]) -> None:
    for key, value in values.items():
        target[f"{prefix}_{key}"] += value


def cmd_apply(args: argparse.Namespace) -> None:
    operation = "profile-apply"
    require_docx(operation)
    src = input_docx(args.input, operation)
    dst = output_docx(args.output, operation)
    profile = PROFILES[args.profile]
    features = risky_features(src)
    if (features["tracked_changes"] or features["macros"]) and not args.allow_complex:
        emit({
            "status": "unsupported",
            "operation": operation,
            "reason": "tracked changes or macros require a preservation-safe path before profile formatting",
            "features": features,
            "office_fallback_allowed": False,
        })
    if src != dst:
        shutil.copy2(src, dst)
    doc = Document(str(dst))
    scopes = {part.strip() for part in args.scopes.split(",") if part.strip()} if args.scopes else set(profile["safe_default_scopes"])
    allowed = {"page", "body", "headings", "tables"}
    if not scopes <= allowed:
        fail(f"invalid scopes: {sorted(scopes - allowed)}", operation)
    changed: Counter[str] = Counter()
    if "page" in scopes:
        changed["sections"] += apply_page(doc, profile["page"])
    if "body" in scopes:
        merge_counts(changed, "body", apply_body(doc, profile["body"]))
    if "headings" in scopes:
        merge_counts(changed, "heading", apply_headings(doc, profile.get("headings", {})))
    if "tables" in scopes:
        merge_counts(changed, "table", apply_tables(doc, profile.get("table", profile["body"])))
    if args.structure_spec:
        merge_counts(changed, "structure", apply_structure_spec(doc, args.structure_spec))
        # Structure assignment happens after existing heading materialization;
        # apply heading specs once more to newly assigned paragraphs.
        if profile.get("headings"):
            merge_counts(changed, "heading_after_structure", apply_headings(doc, profile["headings"]))
    doc.save(str(dst))
    emit({
        "status": "ok",
        "operation": operation,
        "profile": args.profile,
        "output": str(dst),
        "scopes": sorted(scopes),
        "default_body": {"east_asia": "宋体", "size_pt": 12.0},
        "changes": dict(changed),
    })


def cmd_profiles(_: argparse.Namespace) -> None:
    emit({
        "status": "ok",
        "operation": "profiles",
        "profiles": {
            name: {
                "default_body": {"east_asia": spec["body"]["east_asia"], "size_pt": spec["body"]["size"]},
                "line_spacing": spec["body"]["line_spacing"],
                "first_indent_pt": spec["body"]["first_indent_pt"],
                "safe_default_scopes": spec["safe_default_scopes"],
                "heading_styles": list(spec.get("headings", {}).keys()),
            }
            for name, spec in PROFILES.items()
        },
    })


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="LegalWork semantic legal DOCX profile worker")
    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("profiles")
    p.set_defaults(func=cmd_profiles)

    p = sub.add_parser("apply")
    p.add_argument("--input", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--profile", choices=sorted(PROFILES), required=True)
    p.add_argument("--scopes", default="")
    p.add_argument("--structure-spec")
    p.add_argument("--allow-complex", action="store_true")
    p.set_defaults(func=cmd_apply)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
