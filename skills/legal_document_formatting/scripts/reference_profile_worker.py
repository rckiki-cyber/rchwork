#!/usr/bin/env python3
"""Extract and apply a compact Word formatting profile from a reference DOCX.

The worker is intentionally deterministic and compact: it never renders HTML
or returns document text. It can inspect a reference document or apply its
page/style/body conventions to another DOCX while preserving target content.
"""
from __future__ import annotations

import argparse
import json
import shutil
from collections import Counter
from pathlib import Path
from typing import Any, Iterable

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except Exception as exc:  # pragma: no cover
    Document = None  # type: ignore[assignment]
    IMPORT_ERROR = str(exc)
else:
    IMPORT_ERROR = ""

SPECIAL_PREFIXES = (
    "title", "subtitle", "heading", "caption", "quote", "footnote",
    "header", "footer", "toc", "目录", "标题", "题注", "引文",
)


def emit(payload: dict[str, Any], code: int = 0) -> None:
    print(json.dumps(payload, ensure_ascii=False, separators=(",", ":")))
    raise SystemExit(code)


def fail(message: str, code: int = 1) -> None:
    emit({"status": "error", "operation": "reference-profile", "error": message}, code)


def require_docx() -> None:
    if Document is None:
        fail(f"python-docx unavailable: {IMPORT_ERROR}")


def docx_path(raw: str) -> Path:
    path = Path(raw).expanduser().resolve()
    if not path.is_file():
        fail(f"file not found: {path}")
    if path.suffix.lower() != ".docx":
        fail(f"reference-driven formatting only supports .docx (got {path.suffix or 'no extension'})")
    return path


def east_asia_font(target: Any) -> str | None:
    try:
        rpr = target._element.get_or_add_rPr()
        fonts = rpr.rFonts
        return None if fonts is None else fonts.get(qn("w:eastAsia"))
    except Exception:
        return None


def emu_to_pt(value: Any) -> float | None:
    try:
        return round(float(value.pt), 2) if value is not None else None
    except Exception:
        return None


def line_spacing_value(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return round(float(value), 3)
    return emu_to_pt(value)


def alignment_value(value: Any) -> int | None:
    try:
        return None if value is None else int(value)
    except Exception:
        return None


def paragraph_format(paragraph_or_style: Any) -> dict[str, Any]:
    fmt = paragraph_or_style.paragraph_format
    return {
        "alignment": alignment_value(fmt.alignment),
        "line_spacing": line_spacing_value(fmt.line_spacing),
        "first_line_indent_pt": emu_to_pt(fmt.first_line_indent),
        "left_indent_pt": emu_to_pt(fmt.left_indent),
        "right_indent_pt": emu_to_pt(fmt.right_indent),
        "space_before_pt": emu_to_pt(fmt.space_before),
        "space_after_pt": emu_to_pt(fmt.space_after),
        "keep_with_next": bool(fmt.keep_with_next) if fmt.keep_with_next is not None else None,
    }


def font_format(run_or_style: Any) -> dict[str, Any]:
    font = run_or_style.font
    return {
        "east_asia": east_asia_font(run_or_style),
        "latin": font.name,
        "size_pt": emu_to_pt(font.size),
        "bold": bool(font.bold) if font.bold is not None else None,
        "italic": bool(font.italic) if font.italic is not None else None,
    }


def merge_font(primary: dict[str, Any], fallback: dict[str, Any]) -> dict[str, Any]:
    return {key: primary.get(key) if primary.get(key) is not None else fallback.get(key) for key in primary}


def merge_paragraph(primary: dict[str, Any], fallback: dict[str, Any]) -> dict[str, Any]:
    return {key: primary.get(key) if primary.get(key) is not None else fallback.get(key) for key in primary}


def is_body(paragraph: Any) -> bool:
    name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
    return bool(paragraph.text.strip()) and not any(name.startswith(prefix) for prefix in SPECIAL_PREFIXES)


def weighted_mode(values: Iterable[tuple[Any, int]]) -> Any | None:
    counts: Counter[Any] = Counter()
    for value, weight in values:
        if value is not None and weight > 0:
            counts[value] += weight
    return counts.most_common(1)[0][0] if counts else None


def dominant_body_format(doc: Any) -> dict[str, Any]:
    normal = doc.styles["Normal"]
    normal_font = font_format(normal)
    normal_para = paragraph_format(normal)
    latin: list[tuple[str | None, int]] = []
    east_asia: list[tuple[str | None, int]] = []
    sizes: list[tuple[float | None, int]] = []
    alignments: list[tuple[int | None, int]] = []
    line_spacings: list[tuple[float | None, int]] = []
    first_indents: list[tuple[float | None, int]] = []
    spaces_before: list[tuple[float | None, int]] = []
    spaces_after: list[tuple[float | None, int]] = []

    for paragraph in doc.paragraphs:
        if not is_body(paragraph):
            continue
        weight = max(1, len(paragraph.text.strip()))
        pf = paragraph_format(paragraph)
        alignments.append((pf["alignment"], weight))
        line_spacings.append((pf["line_spacing"], weight))
        first_indents.append((pf["first_line_indent_pt"], weight))
        spaces_before.append((pf["space_before_pt"], weight))
        spaces_after.append((pf["space_after_pt"], weight))
        style_font = font_format(paragraph.style)
        for run in paragraph.runs:
            run_weight = max(1, len(run.text.strip()))
            rf = merge_font(font_format(run), style_font)
            latin.append((rf["latin"], run_weight))
            east_asia.append((rf["east_asia"], run_weight))
            sizes.append((rf["size_pt"], run_weight))

    font = {
        "latin": weighted_mode(latin) or normal_font["latin"],
        "east_asia": weighted_mode(east_asia) or normal_font["east_asia"],
        "size_pt": weighted_mode(sizes) or normal_font["size_pt"],
        "bold": normal_font["bold"],
        "italic": normal_font["italic"],
    }
    para = {
        "alignment": weighted_mode(alignments) if alignments else normal_para["alignment"],
        "line_spacing": weighted_mode(line_spacings) if line_spacings else normal_para["line_spacing"],
        "first_line_indent_pt": weighted_mode(first_indents) if first_indents else normal_para["first_line_indent_pt"],
        "left_indent_pt": normal_para["left_indent_pt"],
        "right_indent_pt": normal_para["right_indent_pt"],
        "space_before_pt": weighted_mode(spaces_before) if spaces_before else normal_para["space_before_pt"],
        "space_after_pt": weighted_mode(spaces_after) if spaces_after else normal_para["space_after_pt"],
        "keep_with_next": normal_para["keep_with_next"],
    }
    return {"font": font, "paragraph": para}


def style_profile(doc: Any, style_name: str) -> dict[str, Any] | None:
    try:
        style = doc.styles[style_name]
    except KeyError:
        return None
    return {
        "font": font_format(style),
        "paragraph": paragraph_format(style),
    }


def page_profile(section: Any) -> dict[str, float | None]:
    def cm(value: Any) -> float | None:
        try:
            return round(float(value.cm), 3) if value is not None else None
        except Exception:
            return None

    return {
        "width_cm": cm(section.page_width),
        "height_cm": cm(section.page_height),
        "top_cm": cm(section.top_margin),
        "bottom_cm": cm(section.bottom_margin),
        "left_cm": cm(section.left_margin),
        "right_cm": cm(section.right_margin),
        "header_cm": cm(section.header_distance),
        "footer_cm": cm(section.footer_distance),
    }


def has_page_field(doc: Any) -> bool:
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            xml = paragraph._p.xml
            if "PAGE" in xml or "NUMPAGES" in xml:
                return True
    return False


def extract_profile(doc: Any, source: Path) -> dict[str, Any]:
    profile = {
        "source": str(source),
        "page": page_profile(doc.sections[0]),
        "body": dominant_body_format(doc),
        "styles": {},
        "footer": {"has_page_number": has_page_field(doc)},
    }
    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        spec = style_profile(doc, style_name)
        if spec:
            profile["styles"][style_name] = spec
    return profile


def set_font(target: Any, spec: dict[str, Any]) -> None:
    latin = spec.get("latin")
    east_asia = spec.get("east_asia")
    size = spec.get("size_pt")
    if latin:
        target.font.name = latin
    if size:
        target.font.size = Pt(float(size))
    if spec.get("bold") is not None:
        target.font.bold = bool(spec["bold"])
    if spec.get("italic") is not None:
        target.font.italic = bool(spec["italic"])
    if latin or east_asia:
        fonts = target._element.get_or_add_rPr().get_or_add_rFonts()
        if latin:
            fonts.set(qn("w:ascii"), latin)
            fonts.set(qn("w:hAnsi"), latin)
            fonts.set(qn("w:cs"), latin)
        if east_asia:
            fonts.set(qn("w:eastAsia"), east_asia)


def apply_paragraph(target: Any, spec: dict[str, Any]) -> None:
    fmt = target.paragraph_format
    if spec.get("alignment") is not None:
        try:
            fmt.alignment = WD_ALIGN_PARAGRAPH(int(spec["alignment"]))
        except Exception:
            pass
    value = spec.get("line_spacing")
    if value is not None:
        # Small values are Word line-spacing multipliers; larger values are points.
        fmt.line_spacing = float(value) if float(value) <= 4 else Pt(float(value))
    for key, attr in (
        ("first_line_indent_pt", "first_line_indent"),
        ("left_indent_pt", "left_indent"),
        ("right_indent_pt", "right_indent"),
        ("space_before_pt", "space_before"),
        ("space_after_pt", "space_after"),
    ):
        if spec.get(key) is not None:
            setattr(fmt, attr, Pt(float(spec[key])))
    if spec.get("keep_with_next") is not None:
        fmt.keep_with_next = bool(spec["keep_with_next"])


def apply_page(target_doc: Any, profile: dict[str, Any]) -> None:
    page = profile["page"]
    for section in target_doc.sections:
        if page.get("width_cm"):
            section.page_width = Cm(float(page["width_cm"]))
        if page.get("height_cm"):
            section.page_height = Cm(float(page["height_cm"]))
        for key, attr in (
            ("top_cm", "top_margin"),
            ("bottom_cm", "bottom_margin"),
            ("left_cm", "left_margin"),
            ("right_cm", "right_margin"),
            ("header_cm", "header_distance"),
            ("footer_cm", "footer_distance"),
        ):
            if page.get(key) is not None:
                setattr(section, attr, Cm(float(page[key])))


def apply_profile(target_doc: Any, profile: dict[str, Any], materialize_body: bool) -> dict[str, int]:
    changed = Counter()
    apply_page(target_doc, profile)
    changed["sections"] += len(target_doc.sections)

    body = profile["body"]
    normal = target_doc.styles["Normal"]
    set_font(normal, body["font"])
    apply_paragraph(normal, body["paragraph"])
    changed["styles"] += 1

    for style_name, spec in profile.get("styles", {}).items():
        try:
            style = target_doc.styles[style_name]
        except KeyError:
            continue
        set_font(style, spec["font"])
        apply_paragraph(style, spec["paragraph"])
        changed["styles"] += 1

    if materialize_body:
        for paragraph in target_doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style is not None else ""
            if is_body(paragraph):
                apply_paragraph(paragraph, body["paragraph"])
                for run in paragraph.runs:
                    set_font(run, body["font"])
                    changed["body_runs"] += 1
                changed["body_paragraphs"] += 1
            elif style_name in profile.get("styles", {}):
                spec = profile["styles"][style_name]
                apply_paragraph(paragraph, spec["paragraph"])
                for run in paragraph.runs:
                    set_font(run, spec["font"])
                    changed["heading_runs"] += 1
                changed["headings"] += 1
        for table in target_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_font(run, body["font"])
                            changed["table_runs"] += 1
    return dict(changed)


def cmd_inspect(args: argparse.Namespace) -> None:
    require_docx()
    reference = docx_path(args.input)
    doc = Document(str(reference))
    emit({
        "status": "ok",
        "operation": "reference-inspect",
        "profile": extract_profile(doc, reference),
    })


def cmd_apply(args: argparse.Namespace) -> None:
    require_docx()
    reference = docx_path(args.reference)
    source = docx_path(args.input)
    output = Path(args.output).expanduser().resolve()
    if output.suffix.lower() != ".docx":
        fail("output must end with .docx")
    output.parent.mkdir(parents=True, exist_ok=True)
    if source != output:
        shutil.copy2(source, output)
    reference_doc = Document(str(reference))
    target_doc = Document(str(output))
    profile = extract_profile(reference_doc, reference)
    changed = apply_profile(target_doc, profile, materialize_body=not args.styles_only)
    target_doc.save(str(output))
    emit({
        "status": "ok",
        "operation": "format-like-reference",
        "reference": str(reference),
        "output": str(output),
        "changed": changed,
        "profile_summary": {
            "page": profile["page"],
            "body": profile["body"],
            "styles": list(profile["styles"].keys()),
            "footer": profile["footer"],
        },
        "limitations": [
            "headers/footers content is not copied",
            "custom numbering, fields, comments, and tracked changes are preserved in the target but not cloned from the reference",
        ],
    })


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="LegalWork reference-driven DOCX formatter")
    sub = parser.add_subparsers(dest="command", required=True)

    inspect = sub.add_parser("inspect")
    inspect.add_argument("--input", required=True)
    inspect.set_defaults(func=cmd_inspect)

    apply = sub.add_parser("apply")
    apply.add_argument("--reference", required=True)
    apply.add_argument("--input", required=True)
    apply.add_argument("--output", required=True)
    apply.add_argument("--styles-only", action="store_true")
    apply.set_defaults(func=cmd_apply)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
