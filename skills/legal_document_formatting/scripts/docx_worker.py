#!/usr/bin/env python3
"""Deterministic DOCX worker for LegalWork's legal-document-formatting skill.

The worker emits one compact JSON object and never renders the whole document
into HTML. Unsupported-but-valid requests return exit code 0 with
status="unsupported" plus a one-time Office fallback ticket.
"""
from __future__ import annotations

import argparse
import json
import re
import shutil
import tempfile
import uuid
import zipfile
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

UNSUPPORTED_MARKER = "LEGALWORK_DOCUMENT_UNSUPPORTED"

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except Exception as exc:  # pragma: no cover - environment-dependent
    Document = None  # type: ignore[assignment]
    IMPORT_ERROR = str(exc)
else:
    IMPORT_ERROR = ""

PROFILES: dict[str, dict[str, Any]] = {
    "legal-default": {
        "body_east_asia": "宋体",
        "body_latin": "Times New Roman",
        "body_size": 12.0,
        "line_spacing": 1.5,
        "first_line_chars": 2.0,
        "page": {
            "width_cm": 21.0,
            "height_cm": 29.7,
            "top_cm": 2.54,
            "bottom_cm": 2.54,
            "left_cm": 3.17,
            "right_cm": 3.17,
            "header_cm": 1.5,
            "footer_cm": 1.75,
        },
        "title": {"east_asia": "黑体", "latin": "Times New Roman", "size": 18.0, "bold": True},
        "heading1": {"east_asia": "黑体", "latin": "Times New Roman", "size": 14.0, "bold": True},
        "heading2": {"east_asia": "黑体", "latin": "Times New Roman", "size": 12.0, "bold": True},
        "heading3": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True},
    },
    "academic": {
        "body_east_asia": "宋体",
        "body_latin": "Times New Roman",
        "body_size": 12.0,
        "line_spacing": 1.5,
        "first_line_chars": 2.0,
        "page": {
            "width_cm": 21.0,
            "height_cm": 29.7,
            "top_cm": 2.54,
            "bottom_cm": 2.54,
            "left_cm": 3.17,
            "right_cm": 3.17,
            "header_cm": 1.5,
            "footer_cm": 1.75,
        },
        "title": {"east_asia": "黑体", "latin": "Times New Roman", "size": 18.0, "bold": True},
        "heading1": {"east_asia": "黑体", "latin": "Times New Roman", "size": 15.0, "bold": True},
        "heading2": {"east_asia": "黑体", "latin": "Times New Roman", "size": 14.0, "bold": True},
        "heading3": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True},
    },
    "litigation": {
        "body_east_asia": "宋体",
        "body_latin": "Times New Roman",
        "body_size": 12.0,
        "line_spacing": 1.5,
        "first_line_chars": 2.0,
        "page": {
            "width_cm": 21.0,
            "height_cm": 29.7,
            "top_cm": 2.5,
            "bottom_cm": 2.5,
            "left_cm": 3.0,
            "right_cm": 2.5,
            "header_cm": 1.5,
            "footer_cm": 1.75,
        },
        "title": {"east_asia": "黑体", "latin": "Times New Roman", "size": 18.0, "bold": True},
        "heading1": {"east_asia": "黑体", "latin": "Times New Roman", "size": 14.0, "bold": True},
        "heading2": {"east_asia": "黑体", "latin": "Times New Roman", "size": 12.0, "bold": True},
        "heading3": {"east_asia": "宋体", "latin": "Times New Roman", "size": 12.0, "bold": True},
    },
}

SPECIAL_STYLE_PREFIXES = (
    "title", "subtitle", "heading", "caption", "quote", "footnote",
    "header", "footer", "toc", "目录", "标题", "题注", "引文",
)


def emit(payload: dict[str, Any], exit_code: int = 0) -> None:
    print(json.dumps(payload, ensure_ascii=False, separators=(",", ":")))
    raise SystemExit(exit_code)


def create_fallback_ticket(operation: str, reason: str, detail: Any | None) -> str:
    root = Path(tempfile.gettempdir()) / "legalwork-office-fallback"
    root.mkdir(parents=True, exist_ok=True)
    ticket = root / f"ticket-{uuid.uuid4().hex}.json"
    payload = {
        "marker": UNSUPPORTED_MARKER,
        "status": "unsupported",
        "source": "legal-document-formatting",
        "operation": operation,
        "reason": reason,
        "detail": detail,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    ticket.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    return str(ticket)


def unsupported(reason: str, *, operation: str, detail: Any | None = None) -> None:
    payload: dict[str, Any] = {
        "status": "unsupported",
        "marker": UNSUPPORTED_MARKER,
        "operation": operation,
        "reason": reason,
        "fallback": "office_mcp",
        "fallback_ticket": create_fallback_ticket(operation, reason, detail),
    }
    if detail is not None:
        payload["detail"] = detail
    emit(payload)


def fail(message: str, *, operation: str) -> None:
    emit({"status": "error", "operation": operation, "error": message}, 1)


def require_docx(operation: str) -> None:
    if Document is None:
        unsupported(f"python-docx unavailable: {IMPORT_ERROR}", operation=operation)


def ensure_docx_path(path: str, operation: str, must_exist: bool = True) -> Path:
    p = Path(path).expanduser().resolve()
    if must_exist and not p.is_file():
        fail(f"file not found: {p}", operation=operation)
    if p.suffix.lower() != ".docx":
        unsupported(f"only .docx is handled locally (got {p.suffix or 'no extension'})", operation=operation)
    return p


def copy_for_output(src: Path, output: str, operation: str) -> Path:
    dst = Path(output).expanduser().resolve()
    if dst.suffix.lower() != ".docx":
        fail("output must end with .docx", operation=operation)
    dst.parent.mkdir(parents=True, exist_ok=True)
    if src != dst:
        shutil.copy2(src, dst)
    return dst


def detect_complex_features(path: Path) -> dict[str, bool]:
    features = {
        "tracked_changes": False,
        "comments": False,
        "content_controls": False,
        "fields": False,
        "macros": False,
    }
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
            features["comments"] = "word/comments.xml" in names
            features["macros"] = any(name.lower().endswith("vbaproject.bin") for name in names)
            for name in ("word/document.xml", "word/footnotes.xml", "word/endnotes.xml"):
                if name not in names:
                    continue
                text = zf.read(name).decode("utf-8", errors="ignore")
                if "<w:ins" in text or "<w:del" in text or "<w:moveFrom" in text or "<w:moveTo" in text:
                    features["tracked_changes"] = True
                if "<w:sdt" in text:
                    features["content_controls"] = True
                if "<w:fldSimple" in text or "<w:instrText" in text:
                    features["fields"] = True
    except zipfile.BadZipFile:
        fail("invalid DOCX zip container", operation="inspect")
    return features


def east_asia_font(run_or_style: Any) -> str | None:
    try:
        rpr = run_or_style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        return None if rfonts is None else rfonts.get(qn("w:eastAsia"))
    except Exception:
        return None


def set_font(
    target: Any,
    east_asia: str,
    latin: str,
    size_pt: float,
    bold: bool | None = None,
    *,
    force_black: bool = False,
) -> None:
    target.font.name = latin
    target.font.size = Pt(size_pt)
    if force_black:
        target.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        target.font.bold = bold
    rfonts = target._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:hint"), "eastAsia")


def set_style_font(style: Any, east_asia: str, latin: str, size_pt: float, bold: bool | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        style.font.bold = bold
    rfonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:hint"), "eastAsia")


def ensure_generated_font_table(path: Path) -> None:
    """Declare LegalWork's CJK fonts for stricter DOCX consumers.

    Word resolves ``w:rFonts`` without an explicit font-table entry, but some
    converters treat the missing declaration as an unavailable font and apply
    an arbitrary sans-serif fallback. Keep the user-facing run/style names as
    宋体/黑体 while making the package self-consistent.
    """
    staged = path.with_name(f".{path.name}.font-table-{uuid.uuid4().hex}")
    try:
        with zipfile.ZipFile(path) as source, zipfile.ZipFile(staged, "w") as target:
            for info in source.infolist():
                payload = source.read(info.filename)
                if info.filename == "word/fontTable.xml":
                    text = payload.decode("utf-8", errors="strict")
                    entries: list[str] = []
                    if 'w:name="宋体"' not in text:
                        entries.append(
                            '<w:font w:name="宋体"><w:charset w:val="86"/>'
                            '<w:family w:val="roman"/><w:pitch w:val="variable"/></w:font>'
                        )
                    if 'w:name="黑体"' not in text:
                        entries.append(
                            '<w:font w:name="黑体"><w:charset w:val="86"/>'
                            '<w:family w:val="swiss"/><w:pitch w:val="variable"/></w:font>'
                        )
                    if entries:
                        text = text.replace("</w:fonts>", "".join(entries) + "</w:fonts>")
                        payload = text.encode("utf-8")
                target.writestr(info, payload)
        staged.replace(path)
    finally:
        staged.unlink(missing_ok=True)


def remove_xml_children(parent: Any, *qualified_names: str) -> None:
    if parent is None:
        return
    wanted = {qn(name) for name in qualified_names}
    for child in list(parent):
        if child.tag in wanted:
            parent.remove(child)


def normalize_legal_style_geometry(style: Any, name: str) -> None:
    """Remove visual residue from Word's built-in theme styles.

    python-docx starts from the stock Word template. Merely changing the font
    leaves Title/Heading theme colors, title borders and character spacing in
    place, which is inappropriate for monochrome Chinese legal documents.
    """
    ppr = style._element.get_or_add_pPr()
    rpr = style._element.get_or_add_rPr()
    remove_xml_children(ppr, "w:pBdr", "w:shd", "w:tabs")
    remove_xml_children(rpr, "w:spacing", "w:shd", "w:effect", "w:outline", "w:shadow")

    fmt = style.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.keep_with_next = True
    fmt.keep_together = True
    if name == "Title":
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(18)
        fmt.line_spacing = 1.0
    elif name == "Heading 1":
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(12)
        fmt.space_after = Pt(6)
        fmt.line_spacing = 1.25
    elif name == "Heading 2":
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(9)
        fmt.space_after = Pt(3)
        fmt.line_spacing = 1.25
    elif name == "Heading 3":
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.25


def is_body_paragraph(paragraph: Any) -> bool:
    name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
    return not any(name.startswith(prefix) for prefix in SPECIAL_STYLE_PREFIXES)


def is_list_paragraph(paragraph: Any) -> bool:
    name = (paragraph.style.name if paragraph.style is not None else "").strip().lower()
    ppr = paragraph._p.pPr
    return name.startswith("list") or (ppr is not None and ppr.numPr is not None)


def apply_academic_reference_layout(doc: Any, body_size_pt: float) -> int:
    """Give numbered bibliography entries a real hanging indent.

    Treating references as ordinary body paragraphs produces the inverse
    layout: the marker is indented while wrapped lines return to the margin.
    """
    in_references = False
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading") or style_name == "Title":
            in_references = bool(re.fullmatch(r"(?:[一二三四五六七八九十]+、\s*)?参考文献", text))
            continue
        if not in_references or not re.match(r"^\[\d+\]", text):
            continue
        fmt = paragraph.paragraph_format
        fmt.left_indent = Pt(body_size_pt * 2)
        fmt.right_indent = Pt(0)
        fmt.first_line_indent = Pt(-body_size_pt * 2)
        changed += 1
    return changed


def apply_page(section: Any, page: dict[str, float]) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(page["width_cm"])
    section.page_height = Cm(page["height_cm"])
    section.top_margin = Cm(page["top_cm"])
    section.bottom_margin = Cm(page["bottom_cm"])
    section.left_margin = Cm(page["left_cm"])
    section.right_margin = Cm(page["right_cm"])
    section.header_distance = Cm(page["header_cm"])
    section.footer_distance = Cm(page["footer_cm"])


def ensure_style(doc: Any, name: str, style_type: Any = WD_STYLE_TYPE.PARAGRAPH) -> Any:
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def apply_profile(doc: Any, profile_name: str, scopes: set[str]) -> dict[str, int]:
    profile = PROFILES.get(profile_name)
    if profile is None:
        unsupported(f"unknown profile: {profile_name}", operation="normalize", detail={"profiles": sorted(PROFILES)})
    changed = Counter()
    if "page" in scopes:
        for section in doc.sections:
            apply_page(section, profile["page"])
            changed["sections"] += 1

    if "styles" in scopes or "body" in scopes:
        normal = ensure_style(doc, "Normal")
        set_style_font(normal, profile["body_east_asia"], profile["body_latin"], profile["body_size"])
        pf = normal.paragraph_format
        pf.line_spacing = profile["line_spacing"]
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Pt(profile["body_size"] * profile["first_line_chars"])
        normal._element.get_or_add_pPr().append(OxmlElement("w:widowControl"))
        changed["styles"] += 1

    style_map = {
        "Title": profile["title"],
        "Heading 1": profile["heading1"],
        "Heading 2": profile["heading2"],
        "Heading 3": profile["heading3"],
    }
    if "styles" in scopes or "headings" in scopes:
        for name, cfg in style_map.items():
            style = ensure_style(doc, name)
            set_style_font(style, cfg["east_asia"], cfg["latin"], cfg["size"], cfg["bold"])
            normalize_legal_style_geometry(style, name)
            changed["styles"] += 1

    if "body" in scopes:
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip() or not is_body_paragraph(paragraph):
                continue
            fmt = paragraph.paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.line_spacing = profile["line_spacing"]
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Pt(0)
            fmt.right_indent = Pt(0)
            fmt.first_line_indent = Pt(0) if is_list_paragraph(paragraph) else Pt(
                profile["body_size"] * profile["first_line_chars"]
            )
            for run in paragraph.runs:
                set_font(run, profile["body_east_asia"], profile["body_latin"], profile["body_size"])
                changed["runs"] += 1
            changed["paragraphs"] += 1
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            set_font(run, profile["body_east_asia"], profile["body_latin"], profile["body_size"])
                            changed["table_runs"] += 1
            changed["tables"] += 1

    if "headings" in scopes:
        for paragraph in doc.paragraphs:
            style_name = paragraph.style.name if paragraph.style is not None else ""
            cfg = style_map.get(style_name)
            if not cfg:
                continue
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.first_line_indent = Pt(0)
            if style_name == "Title":
                paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(
                    run,
                    cfg["east_asia"],
                    cfg["latin"],
                    cfg["size"],
                    cfg["bold"],
                    force_black=True,
                )
                changed["heading_runs"] += 1
            changed["headings"] += 1
    if profile_name == "academic" and "body" in scopes:
        changed["reference_paragraphs"] += apply_academic_reference_layout(doc, profile["body_size"])
    return dict(changed)


def summarize(doc: Any, path: Path) -> dict[str, Any]:
    style_counts = Counter((p.style.name if p.style is not None else "(none)") for p in doc.paragraphs if p.text.strip())
    sections = []
    for section in doc.sections[:8]:
        sections.append({
            "page_cm": [round(section.page_width.cm, 2), round(section.page_height.cm, 2)],
            "margins_cm": [
                round(section.top_margin.cm, 2), round(section.bottom_margin.cm, 2),
                round(section.left_margin.cm, 2), round(section.right_margin.cm, 2),
            ],
        })
    return {
        "file": str(path),
        "paragraphs": len(doc.paragraphs),
        "nonempty_paragraphs": sum(1 for p in doc.paragraphs if p.text.strip()),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "top_styles": style_counts.most_common(8),
        "section_sample": sections,
        "features": detect_complex_features(path),
    }


def audit_profile(doc: Any, profile_name: str) -> dict[str, Any]:
    profile = PROFILES[profile_name]
    mismatches = Counter()
    checked = 0
    samples: list[dict[str, Any]] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip() or not is_body_paragraph(paragraph):
            continue
        checked += 1
        for run in paragraph.runs:
            if not run.text.strip():
                continue
            latin = run.font.name or paragraph.style.font.name
            ea = east_asia_font(run) or east_asia_font(paragraph.style)
            size = run.font.size.pt if run.font.size else (paragraph.style.font.size.pt if paragraph.style.font.size else None)
            bad = []
            if latin and latin != profile["body_latin"]:
                mismatches["latin_font"] += 1
                bad.append(f"latin={latin}")
            if ea and ea != profile["body_east_asia"]:
                mismatches["east_asia_font"] += 1
                bad.append(f"eastAsia={ea}")
            if size is not None and abs(size - profile["body_size"]) > 0.05:
                mismatches["font_size"] += 1
                bad.append(f"size={size}")
            if bad and len(samples) < 5:
                samples.append({"paragraph": idx + 1, "issues": bad, "text": paragraph.text[:80]})
    return {"checked_body_paragraphs": checked, "mismatches": dict(mismatches), "samples": samples}


def cmd_inspect(args: argparse.Namespace) -> None:
    operation = "inspect"
    require_docx(operation)
    path = ensure_docx_path(args.input, operation)
    try:
        emit({"status": "ok", "operation": operation, "summary": summarize(Document(str(path)), path)})
    except SystemExit:
        raise
    except Exception as exc:
        fail(str(exc), operation=operation)


def cmd_normalize(args: argparse.Namespace) -> None:
    operation = "normalize"
    require_docx(operation)
    src = ensure_docx_path(args.input, operation)
    features = detect_complex_features(src)
    if (features["tracked_changes"] or features["macros"]) and not args.allow_complex:
        unsupported(
            "document contains tracked changes or macros; local normalization would not guarantee full preservation",
            operation=operation,
            detail=features,
        )
    dst = copy_for_output(src, args.output, operation)
    scopes = {part.strip() for part in args.scopes.split(",") if part.strip()}
    allowed = {"page", "styles", "body", "headings"}
    if not scopes or not scopes <= allowed:
        fail(f"invalid scopes: {sorted(scopes - allowed)}", operation=operation)
    try:
        doc = Document(str(dst))
        changes = apply_profile(doc, args.profile, scopes)
        doc.save(str(dst))
        reloaded = Document(str(dst))
        audit = audit_profile(reloaded, args.profile) if "body" in scopes else {"checked_body_paragraphs": 0, "mismatches": {}}
        emit({
            "status": "ok",
            "operation": operation,
            "output": str(dst),
            "profile": args.profile,
            "scopes": sorted(scopes),
            "changes": changes,
            "audit": audit,
        })
    except SystemExit:
        raise
    except Exception as exc:
        fail(str(exc), operation=operation)


def cmd_page(args: argparse.Namespace) -> None:
    operation = "page"
    require_docx(operation)
    src = ensure_docx_path(args.input, operation)
    dst = copy_for_output(src, args.output, operation)
    try:
        doc = Document(str(dst))
        base = dict(PROFILES[args.profile]["page"])
        overrides = {
            "top_cm": args.top,
            "bottom_cm": args.bottom,
            "left_cm": args.left,
            "right_cm": args.right,
            "header_cm": args.header,
            "footer_cm": args.footer,
        }
        for key, value in overrides.items():
            if value is not None:
                base[key] = float(value)
        for section in doc.sections:
            apply_page(section, base)
        doc.save(str(dst))
        emit({"status": "ok", "operation": operation, "output": str(dst), "sections": len(doc.sections), "page": base})
    except Exception as exc:
        fail(str(exc), operation=operation)


def replace_in_container(paragraphs: Iterable[Any], old: str, new: str) -> tuple[int, int]:
    replaced = 0
    spanning = 0
    for paragraph in paragraphs:
        if old not in paragraph.text:
            continue
        local = 0
        for run in paragraph.runs:
            count = run.text.count(old)
            if count:
                run.text = run.text.replace(old, new)
                local += count
        if local:
            replaced += local
        elif old in paragraph.text:
            spanning += 1
    return replaced, spanning


def all_paragraphs(doc: Any) -> Iterable[Any]:
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def cmd_replace(args: argparse.Namespace) -> None:
    operation = "replace"
    require_docx(operation)
    src = ensure_docx_path(args.input, operation)
    dst = copy_for_output(src, args.output, operation)
    try:
        doc = Document(str(dst))
        replaced, spanning = replace_in_container(all_paragraphs(doc), args.old, args.new)
        if spanning and not args.allow_reflow:
            unsupported(
                "target text crosses Word run boundaries; safe local replacement would not preserve formatting",
                operation=operation,
                detail={"spanning_paragraphs": spanning, "single_run_replacements": replaced},
            )
        if spanning and args.allow_reflow:
            for paragraph in all_paragraphs(doc):
                if args.old in paragraph.text:
                    text = paragraph.text.replace(args.old, args.new)
                    paragraph.clear()
                    paragraph.add_run(text)
                    replaced += 1
        if replaced:
            doc.save(str(dst))
        emit({"status": "ok", "operation": operation, "output": str(dst), "replacements": replaced})
    except SystemExit:
        raise
    except Exception as exc:
        fail(str(exc), operation=operation)


INLINE_MARKDOWN_RE = re.compile(
    r"\[([^\]]+)\]\((https?://[^)\s]+)\)"
    r"|\*\*([^*\n]+)\*\*"
    r"|__([^_\n]+)__"
    r"|`([^`\n]+)`"
    r"|(?<!\*)\*([^*\n]+)\*(?!\*)"
    r"|(?<!_)_([^_\n]+)_(?!_)"
)


def split_markdown_table_row(line: str) -> list[str]:
    source = line.strip().strip("|")
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for character in source:
        if escaped:
            current.append(character)
            escaped = False
        elif character == "\\":
            escaped = True
        elif character == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(character)
    if escaped:
        current.append("\\")
    cells.append("".join(current).strip())
    return cells


def is_markdown_table_separator(line: str) -> bool:
    cells = split_markdown_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def collect_footnotes(md_text: str) -> dict[int, str]:
    """从 markdown 原文收集 GFM 脚注定义 `[^N]: 内容`。

    返回 {id: 脚注文本}，仅保留正文实际引用到且定义了的内容。
    """
    definitions: dict[int, str] = {}
    for line in md_text.split("\n"):
        match = re.match(r"^\s*\[\^(\d+)\]\s*[:：]\s*(.+)$", line)
        if match:
            definitions[int(match.group(1))] = match.group(2).strip()
    referenced = set()
    for line in md_text.split("\n"):
        for m in re.finditer(r"\[\^(\d+)\]", line):
            referenced.add(int(m.group(1)))
    return {key: definitions[key] for key in sorted(referenced) if key in definitions}


def _inject_footnote_refs(doc_bytes: bytes, footnotes: dict[int, str]) -> bytes:
    """把 document.xml 里 `[^N]` 文本替换为合法的 Word 脚注引用。

    必须用 XML 解析而非裸字符串替换：直接把 `<w:footnoteReference/>` run 插进
    `<w:t>` 内部会产生非法嵌套，Word 会丢弃/损坏该段及后续内容（实测 4 千字
    正文只剩 1 千）。正确做法是把含 `[^N]` 的 `<w:t>` 拆成
    `<w:t>前文</w:t><w:r>脚注引用</w:r><w:t>后文</w:t>`。
    """
    try:
        from lxml import etree
    except Exception:
        return doc_bytes  # lxml 不可用时保持原样，不冒险破坏
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W = f"{{{W_NS}}}"
    try:
        root = etree.fromstring(doc_bytes)
    except Exception:
        return doc_bytes
    changed = False
    for t in root.iter(f"{W}t"):
        if t.text is None or "[^" not in t.text:
            continue
        parts: list[str] = []
        cursor = 0
        for match in re.finditer(r"\[\^(\d+)\]", t.text):
            fid = int(match.group(1))
            if fid not in footnotes:
                continue
            if match.start() > cursor:
                parts.append(t.text[cursor:match.start()])
            parts.append(f"\x00FN{fid}\x00")
            cursor = match.end()
            changed = True
        if not changed:
            continue
        if cursor < len(t.text):
            parts.append(t.text[cursor:])
        parent = t.getparent()
        if parent is None:
            continue
        # 拆分：文本片段 → w:t；脚注标记 → 独立 w:r（footnoteReference）。
        for part in parts:
            if part.startswith("\x00FN") and part.endswith("\x00"):
                fid = int(part[3:-1])
                run = etree.SubElement(parent, f"{W}r")
                rpr = etree.SubElement(run, f"{W}rPr")
                rstyle = etree.SubElement(rpr, f"{W}rStyle")
                rstyle.set(f"{W}val", "FootnoteReference")
                va = etree.SubElement(rpr, f"{W}vertAlign")
                va.set(f"{W}val", "superscript")
                ref = etree.SubElement(run, f"{W}footnoteReference")
                ref.set(f"{W}id", str(fid))
            else:
                if not part:
                    continue
                new_t = etree.SubElement(parent, f"{W}t")
                new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                new_t.text = part
        # 移除原 w:t
        parent.remove(t)
    if not changed:
        return doc_bytes
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def inject_footnotes(path: Path, footnotes: dict[int, str]) -> None:
    """把 GFM 脚注注入 docx：document.xml 中 `[^N]` → 上标 footnoteReference，
    并写入 word/footnotes.xml、word/_rels/document.xml.rels、[Content_Types].xml。
    用 zipfile 重打包（与 ensure_generated_font_table 一致），失败静默返回。
    """
    if not footnotes:
        return
    staged = path.with_name(f".{path.name}.footnotes-{uuid.uuid4().hex}")
    try:
        with zipfile.ZipFile(path) as source, zipfile.ZipFile(staged, "w") as target:
            names = {info.filename for info in source.infolist()}
            for info in source.infolist():
                payload = source.read(info.filename)
                name = info.filename
                if name == "word/document.xml":
                    payload = _inject_footnote_refs(payload, footnotes)
                elif name == "word/settings.xml":
                    s = payload.decode("utf-8")
                    # 明确脚注位置在页面底部，避免 WPS/Word 因缺少声明而把脚注
                    # 回退显示到文档末尾。
                    if "<w:footnotePr>" not in s:
                        s = s.replace(
                            "<w:zoom",
                            '<w:footnotePr><w:pos w:val="pageBottom"/></w:footnotePr><w:zoom',
                            1,
                        )
                    payload = s.encode("utf-8")
                elif name == "word/_rels/document.xml.rels":
                    rels = payload.decode("utf-8")
                    if "footnotes" not in rels:
                        rels = rels.replace(
                            "</Relationships>",
                            '<Relationship Id="rIdFootnotes" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/></Relationships>',
                        )
                    payload = rels.encode("utf-8")
                elif name == "[Content_Types].xml":
                    ct = payload.decode("utf-8")
                    if "footnotes.xml" not in ct:
                        ct = ct.replace(
                            "</Types>",
                            '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/></Types>',
                        )
                    payload = ct.encode("utf-8")
                target.writestr(info, payload)

            if "word/footnotes.xml" not in names:
                items = "".join(
                    '<w:footnote w:id="{fid}"><w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
                    '<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/><w:vertAlign w:val="superscript"/></w:rPr>'
                    '<w:footnoteRef/></w:r><w:r><w:t xml:space="preserve"> {text}</w:t></w:r></w:p></w:footnote>'.format(
                        fid=fid, text=_escape_xml(text)
                    )
                    for fid, text in footnotes.items()
                )
                fn_xml = (
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    '<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
                    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
                    f"{items}</w:footnotes>"
                )
                target.writestr("word/footnotes.xml", fn_xml.encode("utf-8"))
        staged.replace(path)
    except Exception:
        # 脚注注入是 best-effort，任何失败不阻塞主流程。
        if staged.exists():
            staged.unlink()


def _escape_xml(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def plain_inline_markdown(text: str) -> str:
    def replacement(match: re.Match[str]) -> str:
        if match.group(1) is not None:
            return match.group(1)
        return next((group for group in match.groups()[2:] if group is not None), "")
    return INLINE_MARKDOWN_RE.sub(replacement, text)


def add_hyperlink(paragraph: Any, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:eastAsia"), "宋体")
    fonts.set(qn("w:cs"), "Times New Roman")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "24")
    run_properties.append(fonts)
    run_properties.append(size)
    run_properties.append(size_cs)
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_markdown_runs(paragraph: Any, text: str) -> None:
    cursor = 0
    for match in INLINE_MARKDOWN_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        if match.group(1) is not None and match.group(2) is not None:
            add_hyperlink(paragraph, match.group(1), match.group(2))
        elif match.group(3) is not None or match.group(4) is not None:
            run = paragraph.add_run(match.group(3) or match.group(4) or "")
            run.bold = True
        elif match.group(5) is not None:
            paragraph.add_run(match.group(5))
        else:
            run = paragraph.add_run(match.group(6) or match.group(7) or "")
            run.italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def parse_markdown_blocks(text: str) -> list[dict[str, Any]]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[dict[str, Any]] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append({"type": "paragraph", "text": " ".join(part.strip() for part in paragraph).strip()})
            paragraph.clear()

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            flush(); i += 1; continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            heading_text = heading.group(2).strip()
            # 论文引注采用当页页下脚注，文末"参考文献"清单是多余的（agent 常
            # 同时写脚注定义和文末列表）。遇到参考文献标题就跳过其后的所有
            # 内容直到下一个标题或文件结束，只保留页下脚注。
            if re.match(r"^(?:参考)?文献(?:清单|列表)?$", heading_text):
                flush()
                i += 1
                while i < len(lines):
                    if re.match(r"^#{1,6}\s+", lines[i].strip()):
                        break
                    i += 1
                continue
            flush(); blocks.append({"type": "heading", "level": len(heading.group(1)), "text": heading_text}); i += 1; continue
        if re.match(r"^[-*+]\s+", stripped):
            flush(); blocks.append({"type": "bullet", "text": re.sub(r"^[-*+]\s+", "", stripped)}); i += 1; continue
        if re.match(r"^\d+[.)、]\s*", stripped):
            flush(); blocks.append({"type": "number", "text": re.sub(r"^\d+[.)、]\s*", "", stripped)}); i += 1; continue
        if stripped.startswith(">"):
            flush(); blocks.append({"type": "quote", "text": re.sub(r"^>\s?", "", stripped)}); i += 1; continue
        if re.fullmatch(r"[-*_]{3,}", stripped):
            flush(); i += 1; continue
        if "|" in stripped and i + 1 < len(lines) and is_markdown_table_separator(lines[i + 1]):
            flush()
            rows = [split_markdown_table_row(stripped)]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append(split_markdown_table_row(lines[i]))
                i += 1
            blocks.append({"type": "table", "rows": rows})
            continue
        paragraph.append(stripped)
        i += 1
    flush()
    return blocks


def cmd_from_markdown(args: argparse.Namespace) -> None:
    operation = "from-markdown"
    require_docx(operation)
    md = Path(args.input).expanduser().resolve()
    if not md.is_file():
        fail(f"file not found: {md}", operation=operation)
    out = Path(args.output).expanduser().resolve()
    if out.suffix.lower() != ".docx":
        fail("output must end with .docx", operation=operation)
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc = Document()
        doc.core_properties.author = "LegalWork"
        doc.core_properties.last_modified_by = "LegalWork"
        apply_profile(doc, args.profile, {"page", "styles"})
        blocks = parse_markdown_blocks(md.read_text(encoding="utf-8"))
        for block in blocks:
            kind = block["type"]
            if kind == "heading":
                level = int(block["level"])
                p = doc.add_paragraph(style="Title" if level == 1 else f"Heading {min(3, level - 1)}")
                add_markdown_runs(p, block["text"])
            elif kind == "bullet":
                add_markdown_runs(doc.add_paragraph(style="List Bullet"), block["text"])
            elif kind == "number":
                add_markdown_runs(doc.add_paragraph(style="List Number"), block["text"])
            elif kind == "quote":
                add_markdown_runs(doc.add_paragraph(style="Quote"), block["text"])
            elif kind == "table":
                rows = block["rows"]
                width = max((len(row) for row in rows), default=1)
                table = doc.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        cell = table.cell(r_idx, c_idx)
                        cell.text = plain_inline_markdown(value)
                        if r_idx == 0:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                if rows:
                    header = OxmlElement("w:tblHeader")
                    header.set(qn("w:val"), "true")
                    table.rows[0]._tr.get_or_add_trPr().append(header)
            else:
                add_markdown_runs(doc.add_paragraph(style="Normal"), block["text"])
        apply_profile(doc, args.profile, {"body", "headings"})
        doc.save(str(out))
        # 注入 GFM 页下脚注（[^1]: 定义），把正文 [^N] 渲染为真 Word 脚注。
        inject_footnotes(out, collect_footnotes(md.read_text(encoding="utf-8")))
        ensure_generated_font_table(out)
        reloaded = Document(str(out))
        emit({
            "status": "ok",
            "operation": operation,
            "output": str(out),
            "profile": args.profile,
            "blocks": len(blocks),
            "summary": {"paragraphs": len(reloaded.paragraphs), "tables": len(reloaded.tables)},
            "audit": audit_profile(reloaded, args.profile),
        })
    except SystemExit:
        raise
    except Exception as exc:
        fail(str(exc), operation=operation)


def cmd_template_fill(args: argparse.Namespace) -> None:
    operation = "template-fill"
    require_docx(operation)
    src = ensure_docx_path(args.input, operation)
    dst = copy_for_output(src, args.output, operation)
    try:
        values = json.loads(Path(args.values).read_text(encoding="utf-8"))
        if not isinstance(values, dict):
            fail("values JSON must be an object", operation=operation)
        doc = Document(str(dst))
        unresolved: list[str] = []
        total = 0
        for key, value in values.items():
            old = "{{" + str(key) + "}}"
            replaced, spanning = replace_in_container(all_paragraphs(doc), old, str(value))
            total += replaced
            if spanning:
                unresolved.append(str(key))
        if unresolved:
            unsupported(
                "template placeholders cross Word run boundaries; local replacement cannot guarantee layout preservation",
                operation=operation,
                detail={"keys": unresolved[:20]},
            )
        doc.save(str(dst))
        emit({"status": "ok", "operation": operation, "output": str(dst), "replacements": total})
    except SystemExit:
        raise
    except Exception as exc:
        fail(str(exc), operation=operation)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="LegalWork deterministic DOCX worker")
    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("inspect")
    p.add_argument("--input", required=True)
    p.set_defaults(func=cmd_inspect)

    p = sub.add_parser("normalize")
    p.add_argument("--input", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--profile", choices=sorted(PROFILES), default="legal-default")
    p.add_argument("--scopes", default="page,styles,body,headings")
    p.add_argument("--allow-complex", action="store_true")
    p.set_defaults(func=cmd_normalize)

    p = sub.add_parser("page")
    p.add_argument("--input", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--profile", choices=sorted(PROFILES), default="legal-default")
    p.add_argument("--top", type=float)
    p.add_argument("--bottom", type=float)
    p.add_argument("--left", type=float)
    p.add_argument("--right", type=float)
    p.add_argument("--header", type=float)
    p.add_argument("--footer", type=float)
    p.set_defaults(func=cmd_page)

    p = sub.add_parser("replace")
    p.add_argument("--input", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--old", required=True)
    p.add_argument("--new", required=True)
    p.add_argument("--allow-reflow", action="store_true")
    p.set_defaults(func=cmd_replace)

    p = sub.add_parser("from-markdown")
    p.add_argument("--input", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--profile", choices=sorted(PROFILES), default="legal-default")
    p.set_defaults(func=cmd_from_markdown)

    p = sub.add_parser("template-fill")
    p.add_argument("--input", required=True)
    p.add_argument("--output", required=True)
    p.add_argument("--values", required=True)
    p.set_defaults(func=cmd_template_fill)
    return parser


def main() -> None:
    args = build_parser().parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
