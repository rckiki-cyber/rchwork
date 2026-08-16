from __future__ import annotations

import importlib.util
import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest import mock

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


WORKER = Path(__file__).resolve().parents[1] / "scripts" / "docx_worker.py"
PROFILE_WORKER = Path(__file__).resolve().parents[1] / "scripts" / "legal_profile_worker.py"


def _load_docx_worker_module():
    """Import scripts/docx_worker.py in-process (not via subprocess)."""
    spec = importlib.util.spec_from_file_location("docx_worker_under_test", WORKER)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class FromMarkdownFormattingTest(unittest.TestCase):
    def test_academic_output_does_not_inherit_blue_title_theme(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            source = root / "review.md"
            output = root / "review.docx"
            source.write_text(
                "# 自动化行政行为程序要件重构\n\n"
                "## 摘要\n\n"
                "这是用于检查宋体小四、黑色标题和段落格式的正文。\n\n"
                "### 一、研究进展\n\n"
                "这是第二段正文。\n\n"
                "## 六、参考文献\n\n"
                "[1] 张三. 一个长度足以换行的参考文献标题[J]. 法学研究, 2026(1): 1-20。",
                encoding="utf-8",
            )

            completed = subprocess.run(
                [
                    sys.executable,
                    str(WORKER),
                    "from-markdown",
                    "--input",
                    str(source),
                    "--output",
                    str(output),
                    "--profile",
                    "academic",
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            payload = json.loads(completed.stdout)
            self.assertEqual(payload["status"], "ok")

            with zipfile.ZipFile(output) as archive:
                font_table = archive.read("word/fontTable.xml").decode("utf-8")
            self.assertIn('w:name="宋体"', font_table)
            self.assertIn('w:name="黑体"', font_table)

            doc = Document(output)
            self.assertEqual(doc.core_properties.author, "LegalWork")
            self.assertEqual([p.text for p in doc.paragraphs[:3]], [
                "自动化行政行为程序要件重构",
                "摘要",
                "这是用于检查宋体小四、黑色标题和段落格式的正文。",
            ])

            title = doc.styles["Title"]
            heading1 = doc.styles["Heading 1"]
            heading2 = doc.styles["Heading 2"]
            for style in (title, heading1, heading2):
                self.assertEqual(str(style.font.color.rgb), "000000")
                self.assertIsNone(style._element.get_or_add_pPr().find(qn("w:pBdr")))
                self.assertIsNone(style._element.get_or_add_rPr().find(qn("w:spacing")))

            self.assertEqual(title.paragraph_format.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            self.assertAlmostEqual(title.paragraph_format.space_after.pt, 18.0)
            self.assertAlmostEqual(heading1.paragraph_format.space_before.pt, 12.0)
            self.assertAlmostEqual(heading1.paragraph_format.space_after.pt, 6.0)
            self.assertAlmostEqual(heading2.paragraph_format.space_before.pt, 9.0)
            self.assertAlmostEqual(heading2.paragraph_format.space_after.pt, 3.0)

            body = doc.paragraphs[2]
            self.assertEqual(body.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
            self.assertAlmostEqual(body.paragraph_format.first_line_indent.pt, 24.0)
            run = body.runs[0]
            fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            self.assertEqual(fonts.get(qn("w:eastAsia")), "宋体")
            self.assertEqual(fonts.get(qn("w:hint")), "eastAsia")
            self.assertEqual(str(doc.styles["Normal"].font.color.rgb), "000000")

            reference = next(p for p in doc.paragraphs if p.text.startswith("[1]"))
            self.assertAlmostEqual(reference.paragraph_format.left_indent.pt, 24.0)
            self.assertAlmostEqual(reference.paragraph_format.first_line_indent.pt, -24.0)

    def test_semantic_profile_also_clears_builtin_title_residue(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            source = root / "agreement.docx"
            output = root / "agreement-formatted.docx"
            doc = Document()
            doc.add_paragraph("委托代理协议", style="Title")
            doc.add_paragraph("甲乙双方经协商一致，订立本协议。")
            doc.save(source)

            completed = subprocess.run(
                [
                    sys.executable,
                    str(PROFILE_WORKER),
                    "apply",
                    "--input",
                    str(source),
                    "--output",
                    str(output),
                    "--profile",
                    "engagement-agreement",
                    "--scopes",
                    "body,headings",
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            self.assertEqual(json.loads(completed.stdout)["status"], "ok")
            formatted = Document(output)
            title = formatted.styles["Title"]
            self.assertEqual(str(title.font.color.rgb), "000000")
            self.assertIsNone(title._element.get_or_add_pPr().find(qn("w:pBdr")))
            self.assertIsNone(title._element.get_or_add_rPr().find(qn("w:spacing")))


class FontTableLockDegradationTest(unittest.TestCase):
    """ensure_generated_font_table 必须 best-effort：Windows 上目标文件被锁
    （WinError 5）时静默降级，保留已写好的 docx，绝不让交付失败。"""

    def test_replace_lock_does_not_raise_and_keeps_docx(self) -> None:
        module = _load_docx_worker_module()
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            out = root / "agreement.docx"
            doc = Document()
            doc.add_paragraph("委托代理协议", style="Title")
            doc.save(out)
            before = out.read_bytes()

            real_replace = Path.replace

            def locked_replace(self_path, other, *args, **kwargs):
                raise PermissionError(5, "拒绝访问")

            with mock.patch.object(Path, "replace", locked_replace):
                module.ensure_generated_font_table(out)  # 不应抛异常

            # 原 docx 必须完好保留
            self.assertEqual(out.read_bytes(), before)
            # 临时 staged 文件应被清理
            leftovers = [p for p in root.iterdir() if p.name.startswith(".agreement.docx.font-table-")]
            self.assertEqual(leftovers, [])
            real_replace(out, out)  # noqa: keep reference alive

    def test_replace_lock_transient_then_succeeds(self) -> None:
        module = _load_docx_worker_module()
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            out = root / "agreement.docx"
            doc = Document()
            doc.add_paragraph("委托代理协议", style="Title")
            doc.save(out)

            real_replace = Path.replace
            calls = {"n": 0}

            def flaky_replace(self_path, other, *args, **kwargs):
                calls["n"] += 1
                if calls["n"] <= 2:
                    raise PermissionError(5, "拒绝访问")
                return real_replace(self_path, other, *args, **kwargs)

            with mock.patch.object(Path, "replace", flaky_replace):
                module.ensure_generated_font_table(out)  # 不应抛异常

            self.assertGreaterEqual(calls["n"], 3)
            # 重试成功后 docx 仍是有效 zip 且含字体表声明
            self.assertTrue(zipfile.is_zipfile(out))
            with zipfile.ZipFile(out) as archive:
                font_table = archive.read("word/fontTable.xml").decode("utf-8")
            self.assertIn('w:name="宋体"', font_table)


if __name__ == "__main__":
    unittest.main()
