from __future__ import annotations

import importlib.util
import os
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


WORKER = Path(__file__).resolve().parents[1] / "scripts" / "pdf_worker.py"
SPEC = importlib.util.spec_from_file_location("legalwork_pdf_worker", WORKER)
assert SPEC and SPEC.loader
PDF_WORKER = importlib.util.module_from_spec(SPEC)
SPEC.loader.exec_module(PDF_WORKER)


class PdfFontSafetyTest(unittest.TestCase):
    def test_internal_renderer_embeds_bundled_songti_and_preserves_profile_geometry(self) -> None:
        try:
            bundled_regular, bundled_bold = PDF_WORKER.bundled_cjk_font_paths()
        except RuntimeError as exc:
            self.skipTest(str(exc))
        with tempfile.TemporaryDirectory() as raw_tmp:
            root = Path(raw_tmp)
            source = root / "report.docx"
            output = root / "report.pdf"
            second_output = root / "report-second.pdf"

            doc = Document()
            section = doc.sections[0]
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
            normal = doc.styles["Normal"]
            normal.font.size = Pt(12)
            normal.paragraph_format.line_spacing = 1.5
            normal.paragraph_format.first_line_indent = Pt(24)
            fonts = normal._element.get_or_add_rPr().get_or_add_rFonts()
            fonts.set(qn("w:eastAsia"), "宋体")
            doc.add_paragraph("算法行政的法律治理研究报告", style="Title")
            doc.add_paragraph("这是用于校验宋体小四和一点五倍行距的中文正文。")
            doc.save(source)

            previous_regular = os.environ.get("LEGALWORK_BUNDLED_CJK_FONT_REGULAR")
            previous_bold = os.environ.get("LEGALWORK_BUNDLED_CJK_FONT_BOLD")
            os.environ["LEGALWORK_BUNDLED_CJK_FONT_REGULAR"] = str(bundled_regular)
            os.environ["LEGALWORK_BUNDLED_CJK_FONT_BOLD"] = str(bundled_bold)
            try:
                summary = PDF_WORKER.reportlab_convert(source, output)
                PDF_WORKER.reportlab_convert(source, second_output)
            finally:
                if previous_regular is None:
                    os.environ.pop("LEGALWORK_BUNDLED_CJK_FONT_REGULAR", None)
                else:
                    os.environ["LEGALWORK_BUNDLED_CJK_FONT_REGULAR"] = previous_regular
                if previous_bold is None:
                    os.environ.pop("LEGALWORK_BUNDLED_CJK_FONT_BOLD", None)
                else:
                    os.environ["LEGALWORK_BUNDLED_CJK_FONT_BOLD"] = previous_bold

            self.assertTrue(output.is_file())
            self.assertEqual(summary["body_size_pt"], 12.0)
            self.assertEqual(summary["line_spacing"], 1.5)
            self.assertTrue(PDF_WORKER.docx_requests_songti(source))
            self.assertTrue(PDF_WORKER.pdf_has_songti(output))
            self.assertTrue(PDF_WORKER.pdf_embeds_songti_font_program(output))
            self.assertTrue(any("notoserifsc" in font.lower() for font in PDF_WORKER.pdf_base_fonts(output)))
            self.assertEqual(output.read_bytes(), second_output.read_bytes())

    def test_production_worker_has_no_user_office_discovery_path(self) -> None:
        self.assertFalse(hasattr(PDF_WORKER, "office_candidates"))
        self.assertFalse(hasattr(PDF_WORKER, "try_libreoffice"))
        source = WORKER.read_text(encoding="utf-8")
        self.assertNotIn("LEGALWORK_SOFFICE", source)
        self.assertNotIn("/Applications/LibreOffice", source)
        self.assertNotIn(r"C:\Program Files\LibreOffice", source)

    def test_rejects_substituted_sans_serif_font_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as raw_tmp:
            output = Path(raw_tmp) / "fake.pdf"
            output.write_bytes(
                b"%PDF-1.7\n/Type /Font /BaseFont /AAAAAA+ArialUnicodeMS\n"
                b"/Type /Font /BaseFont /BBBBBB+SIL-Hei-Med-Jian\n%%EOF"
            )
            self.assertFalse(PDF_WORKER.pdf_has_songti(output))


if __name__ == "__main__":
    unittest.main()
